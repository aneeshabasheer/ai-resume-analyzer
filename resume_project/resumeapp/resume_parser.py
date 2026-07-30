import re
import os
import pdfplumber
import docx
import spacy

# Load spaCy NLP model. If it is not downloaded, we will fall back gracefully.
try:
    nlp = spacy.load("en_core_web_sm")
except Exception:
    nlp = None

# A comprehensive list of skills to match against resume text (case-insensitive)


COMMON_SKILLS = [
    # Programming Languages
    'python', 'java', 'c++', 'c#', 'javascript', 'typescript', 'php', 'ruby', 'go', 'rust', 'kotlin', 'swift', 'r', 'sql', 'html', 'css', 'html5', 'css3',
    # Web Frameworks
    'django', 'flask', 'fastapi', 'spring boot', 'express', 'node.js', 'node', 'react', 'react.js', 'angular', 'vue\\.js', 'vue', 'laravel', 'bootstrap', 'jquery', 'next\\.js',
    # Databases
    'postgresql', 'mysql', 'sqlite', 'mongodb', 'redis', 'oracle', 'cassandra',
    # ML / AI / Data Science
    'machine learning', 'deep learning', 'nlp', 'natural language processing', 'computer vision', 'pandas', 'numpy', 'scikit-learn', 'scikit learn', 'tensorflow', 'pytorch', 'keras', 'seaborn', 'matplotlib', 'data analysis', 'data science',
    # Tools / Cloud / DevOps
    'docker', 'kubernetes', 'aws', 'azure', 'gcp', 'git', 'github', 'jenkins', 'ci/cd', 'linux', 'nginx', 'apache',
    # Other domains
    'cybersecurity', 'cloud computing', 'big data', 'blockchain'
]

def extract_text_from_pdf(pdf_path):
    """
    Extracts text from a PDF file using pdfplumber.
    """
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF: {e}")
    return text

def extract_text_from_docx(docx_path):
    """
    Extracts text from a Word document (.docx) using python-docx.
    """
    text = ""
    try:
        doc = docx.Document(docx_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"Error reading DOCX: {e}")
    return text

def extract_name(text):
    """
    Attempts to extract the candidate's name using spaCy NER.
    Falls back to a clean text-heuristic (first non-empty line) if spaCy fails.
    """
    if nlp:
        doc = nlp(text[:1000]) # Scan first 1000 characters for performance
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                # Clean name: remove newlines or extra spaces
                name = ent.text.strip().replace("\n", " ")
                if len(name.split()) >= 2 and len(name.split()) <= 4:
                    return name
    
    # Fallback heuristic: Get the first clean line from text
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    for line in lines[:5]:
        # Filter out lines that look like email, phone, web, or section headers
        if "@" in line or any(char.isdigit() for char in line) or "resume" in line.lower() or "curriculum" in line.lower():
            continue
        if len(line.split()) >= 2 and len(line.split()) <= 4:
            return line
            
    return "Candidate Name"

def extract_email(text):
    """
    Extracts email using a standard regular expression.
    """
    email_pattern = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
    match = re.search(email_pattern, text)
    return match.group(0) if match else "Not Provided"

def extract_phone(text):
    """
    Extracts phone number using standard pattern matching.
    """
    # Hyphen is placed at the end so it is treated literally.
    phone_pattern = r'\+?\d[\d\s()\-]{8,14}\d'

    match = re.search(phone_pattern, text)

    if match:
        return match.group().strip()

    return "Not Provided"

def extract_skills(text):
    skills_found = []
    text_lower = text.lower()

    for skill in COMMON_SKILLS:
        pattern = r'\b' + re.escape(skill) + r'\b'

        if re.search(pattern, text_lower):
            if skill not in skills_found:
                skills_found.append(skill)

    return list(set(skills_found))


def extract_education(text):
    """
    Finds paragraphs mentioning education keywords like Degrees and Universities.
    """
    education_keywords = [
        'bca', 'mca', 'b.tech', 'm.tech', 'bsc', 'msc', 'b.e.', 'm.e.', 'phd',
        'bachelor', 'master', 'degree', 'university', 'college', 'school', 'diploma'
    ]
    edu_lines = []
    lines = text.split("\n")
    for i, line in enumerate(lines):
        line_lower = line.lower()
        if any(keyword in line_lower for keyword in education_keywords):
            # Add current line, plus next line if short, to give context
            context = line.strip()
            if i + 1 < len(lines) and len(lines[i+1].strip()) < 80:
                context += " - " + lines[i+1].strip()
            edu_lines.append(context)
            if len(edu_lines) >= 3: # limit to top 3 entries
                break
    return "\n".join(edu_lines) if edu_lines else "Education Details Not Found"

def extract_experience(text):
    """
    Finds paragraphs containing experience keywords.
    """
    experience_keywords = ['experience', 'work experience', 'internship', 'job', 'employed', 'employment', 'worked at']
    exp_lines = []
    lines = text.split("\n")
    for i, line in enumerate(lines):
        line_lower = line.lower()
        if any(keyword in line_lower for keyword in experience_keywords):
            context = line.strip()
            if i + 1 < len(lines) and len(lines[i+1].strip()) < 80:
                context += " - " + lines[i+1].strip()
            exp_lines.append(context)
            if len(exp_lines) >= 3:
                break
    return "\n".join(exp_lines) if exp_lines else "Fresher / Experience Details Not Found"

def parse_resume_file(file_path):
    """
    Main controller function: determines file type, extracts text,
    and runs analysis helper functions. Returns a structured dictionary.
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        text = extract_text_from_pdf(file_path)
    elif ext == '.docx':
        text = extract_text_from_docx(file_path)
    else:
        text = ""
        
    if not text:
        return {
            'text': '',
            'name': 'Unknown Name',
            'email': 'Not Provided',
            'phone': 'Not Provided',
            'skills': [],
            'education': 'Could not parse text from file.',
            'experience': 'Could not parse text from file.'
        }
        
    name = extract_name(text)
    email = extract_email(text)
    phone = extract_phone(text)
    skills = extract_skills(text)
    education = extract_education(text)
    experience = extract_experience(text)
    
    return {
        'text': text,
        'name': name,
        'email': email,
        'phone': phone,
        'skills': skills,
        'education': education,
        'experience': experience
    }
